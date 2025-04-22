# general libraries
import os
from datetime import datetime
from dotenv import load_dotenv

# email libraries
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from docx import Document

# load environment variables
load_dotenv()


# create document with title, weekday, script, caption
def create_document(title, weekday, script, caption):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_heading("Script", level=2)
    doc.add_paragraph(script)
    doc.add_heading("Caption", level=3)
    doc.add_paragraph(caption)

    filename = f"{weekday}.docx"
    doc.save(filename)
    return filename


# send email with file attachments
def send_email_with_attachments(files):
    msg = MIMEMultipart()
    msg["From"] = os.getenv("SENDER_EMAIL")
    msg["To"] = os.getenv("RECEIVER_EMAIL")
    msg["Subject"] = f"7 Scripts for this Week starting at {datetime.today().date()}"

    for file in files:
        filename = create_document(file['title'], file['weekday'], file['script'], file['caption'])
        with open(f"{file['weekday']}.docx", "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename={filename}")
            msg.attach(part)

    # Send the email
    with smtplib.SMTP("smtp.gmail.com", 587) as server:
        server.starttls()
        server.login(os.getenv("SENDER_EMAIL"), os.getenv("APP_PASSWORD"))
        server.send_message(msg)
        print("Email sent with 7 attachments!")